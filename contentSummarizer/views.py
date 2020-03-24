from django.shortcuts import render
from django.http import HttpResponse
import os
import sys
import time
import docx
from tenancy_summary import tenancy_summary
from live_and_license_summary import live_and_license_summary
import spacy
from django.core.files.storage import FileSystemStorage
from werkzeug.utils import secure_filename
import docx2txt
from pathlib import Path
from django.utils.encoding import smart_str

def index(request):
    return render(request,'index.html')
def contact(request):
    return render(request,'contact.html')
def home(request):
    return render(request,'index.html')
def about(request):
    return render(request,'about.html')
def download(request):
    if request.method == 'POST':
        file = request.FILES['document']
        section = request.POST['legal_section']
        filename = secure_filename(file.name)
        #print(file.name)
        #print(file.size)
        MY_TEXT = docx2txt.process(file)
        print("Section name is : {}".format(section))
    
    
    if(section=='tenancy'):
        with open("tenancy.txt", "w") as text_file:
            print(MY_TEXT, file=text_file)
        fetched_data= open("tenancy.txt","r+")
        lines = fetched_data.readlines()
        a=tenancy_summary(lines)
        doc = docx.Document()
        doc.add_heading(filename, 0)
        doc_para = doc.add_paragraph(a) 
        #doc.save('tenancy_summarization.docx')
        #return render(request,'download.html')
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.doc')
        response['Content-Disposition'] = 'attachment; filename=tenancy_summary.docx'
        doc.save(response)
        return response
# It's usually a good idea to set the 'Content-Length' header too.
# You can also set any other required headers: Cache-Control, etc.
        #return response



    elif(section=='live and licence'):
        with open("leavelicesne.txt", "w") as text_file:
            print(MY_TEXT, file=text_file)
        fetched_data= open("leavelicesne.txt","r+")
        lines = fetched_data.readlines()
        a=live_and_license_summary(lines)
        doc = docx.Document()
        doc.add_heading(filename, 0)
        doc_para = doc.add_paragraph(a)
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.doc')
        response['Content-Disposition'] = 'attachment; filename=leavelicesne.docx'
        doc.save(response)
        return response



    
        
        #a = tenancy_summary(lines)
        #print(a)
        #fs = FileSystemStorage()
        #fs.save(uploaded_file.name,uploaded_file)
    
